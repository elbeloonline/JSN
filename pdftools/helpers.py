
import base64
import os
import tempfile
from xml.dom.minicompat import NodeList
import zipfile
from django.http import HttpResponse
from lxml import etree
from docxtpl import DocxTemplate
from xml.dom import minidom
from datetime import date
from docx import Document
from copy import deepcopy
from docx.shared import RGBColor
import time
from django.conf import settings

def parse_matches(document: Document, context):
    matches = []
    for match in context["SEARCH_MATCH"]:
        match_array = [m.firstChild.nodeValue if m.firstChild is not None else None for m in match]
        if(len(match_array) != 0):
            matches.append(" ".join(match_array))
        else:
            matches.append("Search For Above Name Has Come Back Clear")
    return matches
    
def parse_court_match(document: Document, context):
    matches = []
    for match in context["COURT_MATCHES"]:
        match_array = [m.firstChild.nodeValue if m.firstChild is not None else None for m in match]
        if len(match_array) != 0 and match_array[0] is not None:
            matches.append(" ".join(match_array))
        else:
            matches.append("***Name Clear***")
    return matches

def parse_judgements(document: Document, context):
    matches = []
    for match in context["COURT_MATCHES"]:
        match_array = [m.firstChild.nodeValue if m.firstChild is not None else None for m in match]
        if len(match_array) != 0 and match_array[0] is not None:
            matches.append("**With Judgements** ")
        else:
            matches.append("")
    return matches

def parse_search_type(document: Document, check, context):
    matches = []
    for type in context["SEARCH_TYPES"]:
        for subtype in type:
            if check == 'isState':
                matches.append("Superior Court:" if subtype == 'STATE_DOCKET_LIST' else '')
            elif check == 'isBankruptcy':
                matches.append("Bankruptcy:" if subtype == 'BANKRUPTCY_DOCKET_LIST' else '')
            elif check == 'isUSDC':
                matches.append("United States District Court:" if subtype == 'USDC_DOCKET_LIST' else '')
    return matches

def parse_docket_num(document: Document, context):
    matches = []
    for docket_results in context["DOCKET_NUMS"]:
        for docket_nums in docket_results:
            dns = []
            for docket_num in docket_nums:
                dns.append(f"{docket_num} _____")
            matches.append("\n".join(dns))
    return matches

def parse_date_signed(document: Document, context):
    matches = []
    for item in context["[DATE_SIGNED]"]:
        if item != "":
            matches.append(f"Date signed: {item}")
        else:
            matches.append("")
    return matches    

def parse_state_comments(document: Document, context):
    matches = []
    for item in context["[ALL_COMMENTS]"]:
        if item != "":
            matches.append(f'Comments: {item}')
        else:
            matches.append("")
    return matches

def parse_state_debt_comments(document: Document, context):
    matches = []
    for item in context["[DEBT_COMMENTS]"]:
        if item != "":
            matches.append(f'Note: {item}')
        else:
            matches.append("")
    return matches

def parse_creditor_attorney(document: Document, case_index, creditor_index, context):
    match = ""
    this_creditor = context["CREDITORS"][case_index][creditor_index]
    
    if this_creditor["[ATTY_FIRST]"] != "":
        match = f'{this_creditor["[ATTY_LAST]"]} {this_creditor["[ATTY_FIRST]"]}'
    else:
        match = "Pro Se"
    return match

def parse_debtor_attorney(document: Document, case_index, debtor_index, context):
    match = ""
    this_debtor = context["DEBTORS"][case_index][debtor_index]
    
    if this_debtor["[ATTY_FIRST]"] != "":
        match = f'{this_debtor["[ATTY_LAST]"]} {this_debtor["[ATTY_FIRST]"]}'
    else:
        match = "Pro Se"
    return match

def parse_dob(document: Document, context):
    matches = []
    for item in context["[DOB]"]:
        if item != "":
            matches.append(f', {context["[DOB]"]}')
        else:
            matches.append("")
    return matches

def replace_context(document: Document, template_name, context):
    #Iterate SEARCH_NAME_ELEMENT and copy tables whenever necessary
    if 'SEARCH_NAME_ELEMENT' in context and context["SEARCH_NAME_ELEMENT"] > 1:
        copies = 0
        while copies < context["SEARCH_NAME_ELEMENT"] - 1:
            table_index = 0
            if template_name == 'CoverPage':
                table_index = 1
            table = document.tables[table_index]
            table.add_row()
            row_index = 2
            if template_name == 'PatriotReport':
                row_index = 3
            for row in table.rows:
                new_tr = deepcopy(row._tr)
                table.rows[row_index]._tr.addnext(new_tr)
                row_index += 1
                
            copies += 1
    #Replace matches in table
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.startswith("[MATCHES}))>0"):
                        context["[MATCHES}))>0,True, False)][SEARCH_MATCH})][else]Search For Above Name Has Come Back Clear[:if]"] = parse_matches(document, context)
                    if p.text.startswith("  [ifWithJudg]"):
                        # context["[ifWithJudg]**With Judgments**[:if][ifWithJudg] [COURT_MATCH][courtMatch] [:foreach][else] ***Name Clear***[:if]"] = "***Name Clear***"
                        context["[ifWithJudg]**With Judgments**[:if]"] = parse_judgements(document, context)
                        context["[ifWithJudg] [COURT_MATCH][courtMatch] [:foreach][else] ***Name Clear***[:if]"] = parse_court_match(document, context) 
    
    #Iterate all paragraphs
    paragraphs = list(document.paragraphs)

    docket_paragraphs = []
    #Paragraphs in docket list
    p_index = 0
    copied_paragraphs = []
    for p in paragraphs:
        if 'N_SEARCH_RESULTS' in context and context['N_SEARCH_RESULTS'] > 0 and p_index >= 1 and p.text not in context.keys():
            docket_paragraphs.append(p)
            copies_para = 1
            while copies_para < context['N_SEARCH_RESULTS']:
                output_para = document.add_paragraph()
                for run in p.runs:
                    output_run = output_para.add_run(run.text)
                    # Run's bold data
                    output_run.bold = run.bold
                    # Run's italic data
                    output_run.italic = run.italic
                    # Run's underline data
                    output_run.underline = run.underline
                    # Run's color data
                    output_run.font.color.rgb = run.font.color.rgb
                    # Run's font data
                    output_run.font.name = run.font.name
                    # Run's style
                    output_run.style.name = run.style.name
                # Paragraph's alignment data
                output_para.paragraph_format.alignment = p.paragraph_format.alignment
                copied_paragraphs.append(output_para)
                copies_para += 1
        if 'N_CASES' in context and context['N_CASES'] > 0 and p.text not in context.keys():
            docket_paragraphs.append(p)
            copies_para = 1
            while copies_para < context['N_CASES']:
                output_para = document.add_paragraph()
                for run in p.runs:
                    output_run = output_para.add_run(run.text)
                    # Run's bold data
                    output_run.bold = run.bold
                    # Run's italic data
                    output_run.italic = run.italic
                    # Run's underline data
                    output_run.underline = run.underline
                    # Run's color data
                    output_run.font.color.rgb = run.font.color.rgb
                    # Run's font data
                    output_run.font.name = run.font.name
                    # Run's style
                    output_run.style.name = run.style.name
                # Paragraph's alignment data
                output_para.paragraph_format.alignment = p.paragraph_format.alignment
                copied_paragraphs.append(output_para)
                copies_para += 1
        p_index += 1
    
    docket_paragraphs += copied_paragraphs
    
    if template_name == 'DocketReport':    
        for i in range(len(docket_paragraphs)):
            print(docket_paragraphs[i].text)
            search_type_paragraphs = []
            paragraphs.append(docket_paragraphs[i])
            copies = 0
            while copies < len(context["SEARCH_TYPES"][copies]) - 1:
                if '[isState]' in docket_paragraphs[i].text or '[DOCKET_NUM]' in docket_paragraphs[i].text:
                    output_para = deepcopy(docket_paragraphs[i])
                    docket_paragraphs[i + 2]._p.addnext(output_para._p)
                    search_type_paragraphs.append(output_para)
                copies += 1
            paragraphs += search_type_paragraphs
            
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    print(p.text)
                    paragraphs.append(p)
    
    header = document.sections[0].header
    for t in header.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    print(p.text)
                    paragraphs.append(p)
    for p in header.paragraphs:
        print(p.text)
        paragraphs.append(p)
    
    if template_name == 'DocketReport':
        context["[isState]Superior Court:[:if]"] = parse_search_type(document, 'isState', context)
        context["[isBankruptcy]Bankruptcy:[:if]"] = parse_search_type(document, 'isBankruptcy', context)
        context["[isUSDC]United States District Court:[:if]"] = parse_search_type(document, 'isUSDC', context)
        context["[DOCKET_NUM][${varDocketNum}] _____"] = parse_docket_num(document, context)
   
    for key, value in context.items():
        j = 0
        for p in paragraphs:
            if key in p.text and type(value) is str:
                runs = p.runs
                started = False
                key_index = 0
                found_runs = list()
                found_all = False
                replace_done = False
                for i in range(len(runs)):
                    # case 1: Found in a single run
                    if key in runs[i].text and not started:
                        found_runs.append((i, runs[i].text.find(key), len(key)))
                        text = runs[i].text.replace(key, value)
                        runs[i].text = text
                        replace_done = True
                        found_all = True
                        break
                    
                    if key[key_index] not in runs[i].text and not started:
                        continue
                    
                    # case 2: Search for partial text, find first run
                    if key[key_index] in runs[i].text and runs[i].text[-1] in key and not started:
                        start_index = runs[i].text.find(key[key_index])
                        check_length = len(runs[i].text)
                        for text_index in range(start_index, check_length):
                            if runs[i].text[text_index] != key[key_index]:
                                break
                        if key_index == 0:
                            started = True
                        chars_found = check_length - start_index
                        key_index += chars_found
                        found_runs.append((i, start_index, chars_found))
                        if key_index != len(key):
                            continue
                        else:
                            found_all = True
                            break
                        
                    # case 2: Search for partial text, find subsequent run
                    if key[key_index] in runs[i].text and started and not found_all:
                        chars_found = 0
                        check_length = len(runs[i].text)
                        for text_index in range(0, check_length):
                            if runs[i].text[text_index] == key[key_index]:
                                key_index += 1
                                chars_found += 1
                            else:
                                break
                        found_runs.append((i, 0, chars_found))
                        if key_index == len(key):
                            found_all = True
                            break
                    
                if found_all and not replace_done:
                    for i, item in enumerate(found_runs):
                        index, start, length = [t for t in item]
                        if i == 0:
                            text = runs[index].text.replace(runs[index].text[start : start + length], value)
                            runs[index].text = text
                        else:
                            text = runs[index].text.replace(runs[index].text[start : start + length], '')
                            runs[index].text = text
            elif key in p.text and type(value) is list and type(value[0]) is str:
                runs = p.runs
                started = False
                key_index = 0
                found_runs = list()
                found_all = False
                replace_done = False
                for i in range(len(runs)):
                    if len(runs[i].text) > 0:
                        # case 1: Found in a single run
                        if key == runs[i].text and started == False:
                            found_runs.append((i, runs[i].text.find(key), len(key)))
                            text = runs[i].text.replace(key, value[j])
                            runs[i].text = text
                            replace_done = True
                            found_all = True
                            break
                        if key[key_index] not in runs[i].text and not started:
                            continue
                        
                        # case 2: Search for partial text, find first run
                        if key[key_index] in runs[i].text and runs[i].text[-1] in key and not started:
                            start_index = runs[i].text.find(key[key_index])
                            check_length = len(runs[i].text)
                            for text_index in range(start_index, check_length):
                                if runs[i].text[text_index] != key[key_index]:
                                    break
                            if key_index == 0:
                                started = True
                            chars_found = check_length - start_index
                            key_index += chars_found
                            found_runs.append((i, start_index, chars_found))
                            
                            if key_index != len(key):
                                continue
                            else:
                                found_all = True
                                break
                            
                        # case 2: Search for partial text, find subsequent run
                        if key[key_index] in runs[i].text and started and not found_all:
                            start_index = runs[i].text.find(key[key_index])
                            chars_found = 0
                            check_length = len(runs[i].text)
                            for text_index in range(0, check_length):
                                if runs[i].text[text_index] == key[key_index]:
                                    key_index += 1
                                    chars_found += 1
                                else:
                                    break
                            found_runs.append((i, start_index, chars_found))
                            if key_index == len(key):
                                found_all = True
                                break
                      
                if found_all and not replace_done:
                    for i, item in enumerate(found_runs):
                        index, start, length = [t for t in item]
                        if i == 0:
                            text = runs[index].text.replace(runs[index].text[start : start + length], value[j], 1)
                            runs[index].text = text
                        else:
                            text = runs[index].text.replace(runs[index].text[start : start + length], '', 1)
                            runs[index].text = text
                    j += 1

def replace_docket_common(document: Document, uri, template_name, context):
    newdoc = Document(uri)
    for _ in range(context['N_CASES']-1):
        for element in Document(uri).element.body:
            newdoc.element.body.append(element)
       
    # Duplicate creditor tables if needed
    table_index = 0 
    prev_index = 0
    for table in newdoc.tables:
        copies = 0
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '[CREDITOR]' in paragraph.text or '[CREDITOR_ATTY_FIRST]' in paragraph.text:
                        if table_index == len(context["CREDITORS"]):
                            break
                        if copies == (len(context["CREDITORS"][table_index]) - 1) * 2:
                            prev_index = table_index
                            table_index += 1
                            break
                        new_tr = deepcopy(row._tr)
                        table.rows[-1]._tr.addnext(new_tr)
                        copies += 1
                        if table_index == len(context["CREDITORS"]):
                            break
                        if copies == (len(context["CREDITORS"][table_index]) - 1) * 2:
                            prev_index = table_index
                            table_index += 1
                            break
                    else:
                        break
                if table_index == len(context["CREDITORS"]):
                    break
                if copies == (len(context["CREDITORS"][prev_index]) - 1) * 2:
                    break
            if table_index == len(context["CREDITORS"]):
                break
            if copies == (len(context["CREDITORS"][prev_index]) - 1)  * 2:
                break
    
    # Duplicate debtor tables if needed                    
    table_index = 0 
    prev_index = 0
    for table in newdoc.tables:
        copies = 0 
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '[DEBTOR]' in paragraph.text or '[DEBTOR_STREET_NAME]' in paragraph.text:
                        if table_index == len(context["DEBTORS"]):
                            break
                        if copies == (len(context["DEBTORS"][table_index]) - 1) * 2:
                            prev_index = table_index
                            table_index += 1
                            break
                        new_tr = deepcopy(row._tr)
                        table.rows[-1]._tr.addnext(new_tr)
                        copies += 1
                        if table_index == len(context["DEBTORS"]):
                            break
                        if copies == (len(context["DEBTORS"][table_index]) - 1) * 2:
                            prev_index = table_index
                            table_index += 1
                            break
                    else:
                        break
                if table_index == len(context["DEBTORS"]):
                    break
                if copies == (len(context["DEBTORS"][prev_index]) - 1) * 2:
                    break
            if table_index == len(context["DEBTORS"]):
                break
            if copies == (len(context["DEBTORS"][prev_index]) - 1)  * 2:
                break        

    for section in newdoc.sections:
        header = section.header
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '[CLIENT_REF_NUM]' in paragraph.text:
                            font_name = paragraph.runs[0].font.name if paragraph.runs else None
                            paragraph.text = paragraph.text.replace('[CLIENT_REF_NUM]', context['[CLIENT_REF_NUM]'])
                            for run in paragraph.runs:
                                run.font.name = font_name

    for paragraph in newdoc.paragraphs:
        for key, values in context.items():
            formatted_key = key.replace('[', '').replace(']', '')
            inline = paragraph.runs
            if key.count('[') > 1:
                if key in paragraph.text:
                    font_name = paragraph.runs[0].font.name if paragraph.runs else None
                    paragraph.text = paragraph.text.replace(key, values.pop(0))
                    for run in paragraph.runs:
                        run.font.name = font_name
            else:
                for i in range(len(inline)):
                    if formatted_key in inline[i].text:
                        if isinstance(values, list):
                            inline[i].text = inline[i].text.replace(formatted_key, values.pop(0))
                            inline[i-1].text = inline[i-1].text.replace('[', '')
                            inline[i+1].text = inline[i+1].text.replace(']', '')
                        else:
                            inline[i].text = inline[i].text.replace(formatted_key, values)
                            inline[i - 1].text = inline[i - 1].text.replace('[', '')
                            inline[i + 1].text = inline[i + 1].text.replace(']', '')
                        inline[i].font.color.rgb = RGBColor(0, 0, 0)
                    if inline[i].text == '[' and inline[i+1].text == ':':
                        for _ in range(4):
                            inline[i].text = ''
                            i+=1
    
    creditor_index = 0
    current_creditor_index = 0
    rows_done_creditor = 0
    debtor_index = 0
    current_debtor_index = 0
    rows_done_debtor = 0                        
    for table in newdoc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "[:forEach]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[:forEach]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[case]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[case]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[ifcreditorstreet]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[ifcreditorstreet]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[IF_STREET_PRESENT]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[IF_STREET_PRESENT]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "(No Address)[:if][multipleDebtors]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("(No Address)[:if][multipleDebtors]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[:if][:if]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[:if][:if]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[else][:if]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[else][:if]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[ifaltname][else]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[ifaltname][else]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[DEBT]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[DEBT]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[CREDITOR]" in paragraph.text and template_name != 'CaseReport':
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[CREDITOR]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[DEBTOR]" in paragraph.text and template_name != 'CaseReport':
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[DEBTOR]", '')
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[USDC_CASE_NUMBER]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[USDC_CASE_NUMBER]", context.get("[USDC_CASE_NUMBER]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[ATTORNEY]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[ATTORNEY]", context.get("[ATTORNEY]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[CREDITOR_ATTORNEY]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[CREDITOR_ATTORNEY]", context.get("[CREDITOR_ATTORNEY]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[JUDGMENT_NUMBER]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[JUDGMENT_NUMBER]", context.get("[JUDGMENT_NUMBER]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[DATE_ENTERED]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[DATE_ENTERED]", context.get("[DATE_ENTERED]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[VENUE]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[VENUE]", context.get("[VENUE]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[CASE_NUM]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[CASE_NUM]", context.get("[CASE_NUM]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[ACTION_TYPE]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[ACTION_TYPE]", context.get("[ACTION_TYPE]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[CASE_STATUS]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[CASE_STATUS]", context.get("[CASE_STATUS]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[ifDateSigned]Date Signed: [DATE_SIGNED][:if]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[ifDateSigned]Date Signed: [DATE_SIGNED][:if]", context.get("[ifDateSigned]Date Signed: [DATE_SIGNED][:if]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                    if "[ORIG_DEBT_AMT]" in paragraph.text:
                        font_name = paragraph.runs[0].font.name if paragraph.runs else None
                        paragraph.text = paragraph.text.replace("[ORIG_DEBT_AMT]", context.get("[ORIG_DEBT_AMT]").pop(0))
                        for run in paragraph.runs:
                            run.font.name = font_name
                            
                    if creditor_index < len(context["CREDITORS"]):
                        creditor = context["CREDITORS"][creditor_index]
                        
                        if len(creditor) != 0:
                            this_creditor = creditor[current_creditor_index]
                            if "[CREDITOR][PARTY_NAME]" in paragraph.text:
                                font_name = paragraph.runs[0].font.name if paragraph.runs else None
                                paragraph.text = paragraph.text.replace("[CREDITOR][PARTY_NAME]", this_creditor["[PARTY_NAME]"], 1)
                                for run in paragraph.runs:
                                    run.font.name = font_name
                                rows_done_creditor += 1
                            if "[CREDITOR_STREET_NAME][:if]" in paragraph.text:
                                font_name = paragraph.runs[0].font.name if paragraph.runs else None
                                paragraph.text = paragraph.text.replace("[CREDITOR_STREET_NAME][:if]", this_creditor["[CREDITOR_STREET_NAME]"], 1)
                                for run in paragraph.runs:
                                    run.font.name = font_name
                                rows_done_creditor += 1
                            if "[ifdebtcomment], [DEBT_COMMENT][:if]" in paragraph.text:
                                font_name = paragraph.runs[0].font.name if paragraph.runs else None
                                paragraph.text = paragraph.text.replace("[ifdebtcomment], [DEBT_COMMENT][:if]", this_creditor["[DEBT_COMMENT]"], 1)
                                for run in paragraph.runs:
                                    run.font.name = font_name
                                rows_done_creditor += 1
                            if "[CREDITOR_ATTY_FIRST})>0,True,False)] [CREDITOR_ATTY_LAST] [CREDITOR_ATTY_FIRST][else]Pro Se[:if]" in paragraph.text:
                                font_name = paragraph.runs[0].font.name if paragraph.runs else None
                                paragraph.text = paragraph.text.replace("[CREDITOR_ATTY_FIRST})>0,True,False)] [CREDITOR_ATTY_LAST] [CREDITOR_ATTY_FIRST][else]Pro Se[:if]", parse_creditor_attorney(document, creditor_index, current_creditor_index, context), 1)
                                for run in paragraph.runs:
                                    run.font.name = font_name
                                rows_done_creditor += 1
                                
                            if rows_done_creditor == 4:
                                current_creditor_index += 1
                                rows_done_creditor = 0
                        if current_creditor_index == len(creditor):
                            creditor_index += 1
                            current_creditor_index = 0
                            
                    if debtor_index < len(context["DEBTORS"]):
                        debtor = context["DEBTORS"][debtor_index]
                        
                        if len(debtor) != 0:
                            this_debtor = debtor[current_debtor_index]
                            if "[DEBTOR] [PARTY_NAME]" in paragraph.text:
                                font_name = paragraph.runs[0].font.name if paragraph.runs else None
                                paragraph.text = paragraph.text.replace("[DEBTOR] [PARTY_NAME]", this_debtor["[PARTY_NAME]"], 1)
                                for run in paragraph.runs:
                                    run.font.name = font_name
                                rows_done_debtor += 1
                            if "[DEBTOR_STREET_NAME][else]" in paragraph.text:
                                font_name = paragraph.runs[0].font.name if paragraph.runs else None
                                paragraph.text = paragraph.text.replace("[DEBTOR_STREET_NAME][else]", this_debtor["[DEBTOR_STREET_NAME]"], 1)
                                for run in paragraph.runs:
                                    run.font.name = font_name
                                rows_done_debtor += 1
                            if "[IF_DOB_PRESENT], [DOB][:if]" in paragraph.text:
                                font_name = paragraph.runs[0].font.name if paragraph.runs else None
                                paragraph.text = paragraph.text.replace("[IF_DOB_PRESENT], [DOB][:if]", this_debtor["[DOB]"], 1)
                                for run in paragraph.runs:
                                    run.font.name = font_name
                                rows_done_debtor += 1
                            if "[ifaltname]Alternate Name: [PARTY_ALT][PARTY_NAME]" in paragraph.text:
                                font_name = paragraph.runs[0].font.name if paragraph.runs else None
                                if len(this_debtor["[PARTY_ALT]"]) > current_debtor_index:
                                    paragraph.text = paragraph.text.replace("[ifaltname]Alternate Name: [PARTY_ALT][PARTY_NAME]", this_debtor["[PARTY_ALT]"][current_debtor_index], 1)
                                else:
                                    paragraph.text = paragraph.text.replace("[ifaltname]Alternate Name: [PARTY_ALT][PARTY_NAME]", "", 1)
                                for run in paragraph.runs:
                                    run.font.name = font_name
                                rows_done_debtor += 1
                            if "[DEBTOR_ATTY_FIRST})=NULL(), False, True)][DEBTOR_ATTY_FIRST})>0, True, False)][DEBTOR_ATTY_FIRST] [DEBTOR_ATTY_LAST][else]Pro Se[:if][else]Pro Se[:if]" in paragraph.text:
                                font_name = paragraph.runs[0].font.name if paragraph.runs else None
                                paragraph.text = paragraph.text.replace("[DEBTOR_ATTY_FIRST})=NULL(), False, True)][DEBTOR_ATTY_FIRST})>0, True, False)][DEBTOR_ATTY_FIRST] [DEBTOR_ATTY_LAST][else]Pro Se[:if][else]Pro Se[:if]", parse_debtor_attorney(document, debtor_index, current_debtor_index, context), 1)
                                for run in paragraph.runs:
                                    run.font.name = font_name
                                rows_done_debtor += 1
                            
                            if rows_done_debtor == 5:
                                current_debtor_index += 1
                                rows_done_debtor = 0
                        if current_debtor_index == len(debtor):
                            debtor_index += 1
                            current_debtor_index = 0
                                
    for table in newdoc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:    
                    if paragraph.text == '\t, ' or paragraph.text == '\t':
                        p = paragraph._element
                        p.getparent().remove(p)
                        p._p = p._element = None
                    if paragraph.text == '':
                        inline = paragraph.runs
                        if len(inline)>0:
                            p = paragraph._element
                            p.getparent().remove(p)
                            p._p = p._element = None
                    for run in paragraph.runs:
                        if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                            p = paragraph._element
                            p.getparent().remove(p)
                            p._p = p._element = None
                            
    for paragraph in newdoc.paragraphs:
        if paragraph.text == '\t, ':
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None
        if paragraph.text == '':
            inline = paragraph.runs
            if len(inline)>0:
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None
        for run in paragraph.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None
    
    return newdoc            
            
def replace_patriot(uri, xml_data: minidom.Document):
    client_ref_num = xml_data.getElementsByTagName('CLIENT_REF_NUM')[0].firstChild.nodeValue
    search_to = xml_data.getElementsByTagName('SEARCH_TO')[0].firstChild.nodeValue
    search_name_element = xml_data.getElementsByTagName('SEARCH_NAME_ELEMENT')
    matches = []
    searched_names = []
    search_matches = []
    for i in range(0, len(search_name_element)):
        match = search_name_element[i].getElementsByTagName('MATCHES')[0].childNodes
        searched_name = search_name_element[i].getElementsByTagName('SEARCHED_NAME')[0].firstChild.nodeValue
        search_match = search_name_element[i].getElementsByTagName('SEARCH_MATCH')
        matches.append(match)
        searched_names.append(searched_name)
        search_matches.append(search_match)
    context = {
        '[CLIENT_REF_NUM]': client_ref_num,  
        "[SEARCH_TO]": search_to, 
        'MATCHES': matches, 
        'SEARCH_MATCH': search_matches,
        '[SEARCHED_NAME})]': searched_names, 
        'SEARCH_NAME_ELEMENT': len(search_name_element), 
        "[SEARCH_NAME_ELEMENT]": '', 
        '[:forEach]': '',
    }    
    document = Document(uri)
    replace_context(document, 'PatriotReport', context)
    document.save(os.path.join(".",'jsnetwork_project','media',f'generated_PatriotReport.docx'))
    
def replace_coverpage(uri, xml_data: minidom.Document):
    client_file_num = xml_data.getElementsByTagName('CLIENT_FILE_NUM')[0].firstChild.nodeValue
    name = xml_data.getElementsByTagName('NAME')[0].firstChild.nodeValue
    print(xml_data.getElementsByTagName('ADDRESS')[0].firstChild)
    if xml_data.getElementsByTagName('ADDRESS')[0].firstChild is not None:
        address = xml_data.getElementsByTagName('ADDRESS')[0].firstChild.nodeValue
    else:
        address = "Address not found" 
    city_state_zip = xml_data.getElementsByTagName('CITY_STATE_ZIP')[0].firstChild.nodeValue
    search_name_element = xml_data.getElementsByTagName('SEARCH_NAME_ELEMENT')
    court_matches = []
    search_names = []
    search_froms = []
    search_tos = []
    for i in range(0, len(search_name_element)):
        search_name = search_name_element[i].getElementsByTagName('SEARCH_NAME')[0].firstChild.nodeValue
        court_match = [court for court in search_name_element[i].getElementsByTagName('COURT_MATCH')]
        search_to = search_name_element[i].getElementsByTagName('SEARCH_TO')[0].firstChild.nodeValue
        search_from = search_name_element[i].getElementsByTagName('SEARCH_FROM')[0].firstChild.nodeValue
        search_names.append(search_name)
        court_matches.append(court_match)
        search_tos.append(search_to)
        search_froms.append(search_from)
    now = date.today().strftime('%m/%d/%Y')
    context = {
        '[CLIENT_FILE_NUM]': client_file_num, 
        "[NAME]": name, 
        "[ADDRESS]": address, 
        "[CITY_STATE_ZIP]": city_state_zip, 
        "[SEARCH_NAME]": search_names, 
        "[SEARCH_FROM]": search_froms, 
        "[SEARCH_TO]": search_tos, 
        "[=NOW()]": now, 
        "COURT_MATCHES": court_matches,
        "SEARCH_NAME_ELEMENT": len(search_name_element),
        "[SEARCH_NAME_ELEMENT]": '', 
        '[:forEach]': '',
    }    
    document = Document(uri)
    replace_context(document, 'CoverPage', context)
    document.save(os.path.join(".",'jsnetwork_project','media',f'generated_CoverPage.docx'))
    
def replace_docketreport(uri, xml_data: minidom.Document):
    search_results = xml_data.getElementsByTagName('SEARCH_RESULTS')
    search_result = []
    search_names = []
    search_type_names_all = []
    docket_nums_all = []
    for i in range(0, len(search_results)):
        search_name = search_results[i].getElementsByTagName('SEARCH_NAME')[0].firstChild.nodeValue
        search_type_nodes = search_results[i].getElementsByTagName('SEARCH_TYPE')
        search_type_names = []
        docket_nums_result = []
        for j in range(0, len(search_type_nodes)):
            search_type_name = search_type_nodes[j].getElementsByTagName('SEARCH_TYPE_NAME')[0].firstChild.nodeValue
            docket_num_nodes = search_type_nodes[j].getElementsByTagName('DOCKET_RESULTS')[0].childNodes
            docket_nums = []
            for docket_num_node in docket_num_nodes:
                if(type(docket_num_node) is minidom.Element):
                    docket_num = docket_num_node.firstChild.nodeValue
                    docket_nums.append(docket_num)
            search_type_names.append(search_type_name)
            docket_nums_result.append(docket_nums)
        search_type_names_all.append(search_type_names)
        docket_nums_all.append(docket_nums_result) 
        search_names.append(search_name)
        search_result.append("")
    context = { 
        "N_SEARCH_RESULTS": len(search_results),
        "[SEARCH_RESULTS]": '',
        "[SEARCH_NAME]": search_names,
        "[SEARCH_TYPE]": "",
        "DOCKET_NUMS": docket_nums_all,
        "SEARCH_TYPES": search_type_names_all,
        "[endSearchType]": "", 
        "[endSearchResults]": "", 
        "[endDocketNum]": "",
    }
    document = Document(uri)
    replace_context(document, 'DocketReport', context)
    document.save(os.path.join(".",'jsnetwork_project','media',f'generated_DocketReport.docx'))
    
def replace_usdc(uri, xml_data: minidom.Document):
    client_ref_num = xml_data.getElementsByTagName('CLIENT_REF_NUM')[0].firstChild.nodeValue
    cases = xml_data.getElementsByTagName('case')
    searchnames = []
    usdc_case_numbers = []
    dates_filed = []
    creditors = []
    creditor_attorneys = []
    debtors = []
    attorneys = []
    trustees = []
    dates_discharged = []
    dates_terminated = []
    for i in range(0, len(cases)):
        searchname = cases[i].getElementsByTagName('SEARCHNAME')[0].firstChild.nodeValue
        searchnames.append(searchname)
        usdc_case_number = cases[i].getElementsByTagName('USDC_CASE_NUMBER')[0].firstChild.nodeValue
        usdc_case_numbers.append(usdc_case_number)
        date_filed = cases[i].getElementsByTagName('DATE_FILED')[0].firstChild.nodeValue
        dates_filed.append(date_filed)
        parties = cases[i].getElementsByTagName('PARTIES')
        creditor = get_element_value_from_parties(cases[i], 'CREDITOR')
        creditors.append(creditor)
        creditor_attorney = get_element_value_from_parties(cases[i], 'CREDITOR_ATTORNEY')
        creditor_attorneys.append(creditor_attorney)
        debtor = get_element_value_from_parties(cases[i], 'DEBTOR')
        debtors.append(debtor)
        attorney = get_element_value_from_parties(cases[i], 'ATTORNEY')
        attorneys.append(attorney)
        trustee = get_element_value_from_parties(cases[i], 'TRUSTEE')
        trustees.append(trustee)
        date_discharged = cases[i].getElementsByTagName('DATE_DISCHARGED')[0].firstChild.nodeValue
        dates_discharged.append(date_discharged)
        date_terminated = cases[i].getElementsByTagName('DATE_TERMINATED')[0].firstChild.nodeValue
        dates_terminated.append(date_terminated)
        
    context = {
        "N_CASES": len(cases),
        "[case]": '',
        "[PARTIES]": '',
        "[CLIENT_REF_NUM]": client_ref_num,
        "[USDC_CASE_NUMBER]": usdc_case_numbers,
        "[DATE_FILED]": dates_filed,
        "[CREDITOR]": creditors,
        "[CREDITOR_ATTORNEY]": creditor_attorneys,
        "[DEBTOR]": debtors,
        "[ATTORNEY]": attorneys,
        "[TRUSTEE]": trustees,
        "[DATE_DISCHARGED]": dates_discharged,
        "[DATE_TERMINATED]": dates_terminated,
        "[:forEach]": ''
    }
    document = Document(uri)    
    newdoc = replace_docket_common(document, uri, 'UsdcReport', context)
    newdoc.save(os.path.join(".",'jsnetwork_project','media',f'generated_UsdcReport.docx'))

def replace_bankruptcy(uri, xml_data: minidom.Document):
    client_ref_num = xml_data.getElementsByTagName('CLIENT_REF_NUM')[0].firstChild.nodeValue
    cases = xml_data.getElementsByTagName('case')

    context = {
        "[case]": '',
        "[PARTIES]": '',
        "[CLIENT_REF_NUM]": client_ref_num,
        "[BANKRUPTCY_NUMBER]":[],
        "[FILED]":[],
        "[VOL]": [],
        "[CHAP]": [],
        "[DEBTOR]": [],
        "[SSN]": [],
        "[ADDRESS]": [],
        "[CITYSTATEZIP]": [],
        "[ATTORNEY]": [],
        "[TRUSTEE]": [],
        "[DISCH]": [],
        "[TERM]": [],
        "[:forEach]":''
    }

    for case_element in cases:
        context["[DEBTOR]"].append(get_element_value_from_parties(case_element, 'DEBTOR'))
        context["[SSN]"].append(get_element_value_from_parties(case_element, 'SSN'))
        context["[ADDRESS]"].append(get_element_value_from_parties(case_element, 'ADDRESS'))
        context["[CITYSTATEZIP]"].append(get_element_value_from_parties(case_element, 'CITYSTATEZIP'))
        context["[ATTORNEY]"].append(get_element_value_from_parties(case_element, 'ATTORNEY'))
        context["[TRUSTEE]"].append(get_element_value_from_parties(case_element, 'TRUSTEE'))
        context["[BANKRUPTCY_NUMBER]"].append(get_element_value(case_element, 'BANKRUPTCY_NUMBER'))
        context["[VOL]"].append(get_element_value(case_element, 'VOLUNTARY'))
        context["[CHAP]"].append(get_element_value(case_element, 'CHAPTER'))
        context["[DISCH]"].append(get_element_value(case_element, 'DATE_DISCHARGED'))
        context["[TERM]"].append(get_element_value(case_element, 'DATE_TERMINATED'))
        context["[FILED]"].append(get_element_value(case_element, 'DATE_FILED'))

    context['N_CASES'] = len(context.get('[BANKRUPTCY_NUMBER]'))
    document = Document(uri)
    
    newdoc = replace_docket_common(document, uri, 'BankruptcyReport', context)
    
    temp_word_path  = os.path.join(".",'jsnetwork_project','media', f'generated_BankruptcyReport.docx')
    newdoc.save(temp_word_path)
    
def replace_state(uri, xml_data: minidom.Document):
    client_ref_num = xml_data.getElementsByTagName('CLIENT_REF_NUM')[0].firstChild.nodeValue
    cases = xml_data.getElementsByTagName('case')
    judgment_numbers = []
    dates_entered = []
    venues = []
    case_numbers = []
    dates_signed = []
    action_types = []
    case_statuses = []
    orig_debt_amts = []
    all_creditors = []
    all_debtors = []
    comments = []
    debt_comments = []
    for i in range(0, len(cases)):
        
        judgment_number = cases[i].getElementsByTagName('JUDGMENT_NUMBER')[0].firstChild.nodeValue
        judgment_numbers.append(judgment_number)
        date_entered = cases[i].getElementsByTagName('DATE_ENTERED')[0].firstChild.nodeValue
        dates_entered.append(date_entered)
        venue = cases[i].getElementsByTagName('VENUE')[0].firstChild.nodeValue
        venues.append(venue)
        case_number = cases[i].getElementsByTagName('CASE_NUMBER')[0].firstChild.nodeValue
        case_numbers.append(case_number)
        date_signed = get_element_value(cases[i], 'DATE_SIGNED')
        dates_signed.append(date_signed)
        action_type = cases[i].getElementsByTagName('ACTION_TYPE')[0].firstChild.nodeValue
        action_types.append(action_type)
        case_status = cases[i].getElementsByTagName('CASE_STATUS')[0].firstChild.nodeValue
        case_statuses.append(case_status)
        orig_debt_amt = get_element_values_from_debt(cases[i], 'ORIG_DEBT_AMT')
        orig_debt_amts.append(orig_debt_amt)
        creditors = cases[i].getElementsByTagName('DEBT')[0].getElementsByTagName('CREDITOR')
        debtors = cases[i].getElementsByTagName('DEBT')[0].getElementsByTagName('DEBTOR')
        comment = get_element_value(cases[i], 'COMMENTS')
        comments.append(comment)
        debt_comment = get_element_values_from_debt(cases[i], 'DEBT_COMMENT')
        debt_comments.append(debt_comment)

        debtor_alt_names = []
        
        party_alt = cases[i].getElementsByTagName('DEBT')[0].getElementsByTagName('PARTY_ALT')
        for alt in party_alt:
            alt_party_name = alt.getElementsByTagName('PARTY_NAME') 
            if len(alt_party_name) != 0:
                party_alt_name = alt_party_name[0].firstChild.nodeValue
                debtor_alt_names.append(f"Alternate Name: {party_alt_name}")
            else:
               debtor_alt_names.append("")

        if len(creditors) > 1:
            case_creditors = []
            for c in range(0, len(creditors)):
                creditor_party_name = get_creditor_values_from_debt(cases[i], c, 'PARTY_NAME')
                creditor_street_name = get_creditor_values_from_debt(cases[i], c, 'CREDITOR_STREET_NAME')
                # creditor_debt_comment = get_creditor_values_from_debt(cases[i], c, 'DEBT_COMMENT')
                creditor_debt_comment = debt_comment
                creditor_atty_first = get_creditor_values_from_debt(cases[i], c, 'ATTY_FIRST')
                creditor_atty_last = get_creditor_values_from_debt(cases[i], c, 'ATTY_LAST')
                case_creditors.append({
                    "[PARTY_NAME]": creditor_party_name,
                    "[CREDITOR_STREET_NAME]": creditor_street_name,
                    "[DEBT_COMMENT]": f", {creditor_debt_comment}" if creditor_debt_comment != ""  else "",
                    "[ATTY_FIRST]": creditor_atty_first,
                    "[ATTY_LAST]": creditor_atty_last,
                })
            all_creditors.append(case_creditors)
        else:
            creditor_party_name = get_creditor_values_from_debt(cases[i], 0, 'PARTY_NAME')
            creditor_street_name = get_creditor_values_from_debt(cases[i], 0, 'CREDITOR_STREET_NAME')
            # creditor_debt_comment = get_creditor_values_from_debt(cases[i], 0, 'DEBT_COMMENT')
            creditor_debt_comment = debt_comment
            creditor_atty_first = get_creditor_values_from_debt(cases[i], 0, 'ATTY_FIRST')
            creditor_atty_last = get_creditor_values_from_debt(cases[i], 0, 'ATTY_LAST')
            all_creditors.append([{
                "[PARTY_NAME]": creditor_party_name,
                "[CREDITOR_STREET_NAME]": creditor_street_name,
                "[DEBT_COMMENT]": f", {creditor_debt_comment}" if creditor_debt_comment != ""  else "",
                "[ATTY_FIRST]": creditor_atty_first,
                "[ATTY_LAST]": creditor_atty_last,
            }])

        if len(debtors) > 1:
            case_debtors = []
            for d in range(0, len(debtors)):
                debtor_party_name = get_debtor_values_from_debt(cases[i], d, 'PARTY_NAME')
                debtor_street_name = get_debtor_values_from_debt(cases[i], d, 'DEBTOR_STREET_NAME')
                dob = get_debtor_values_from_debt(cases[i], d, 'DOB')
                debtor_atty_first = get_debtor_values_from_debt(cases[i], d, 'ATTY_FIRST')
                debtor_atty_last = get_debtor_values_from_debt(cases[i], d, 'ATTY_LAST')
                case_debtors.append({
                    "[PARTY_NAME]": debtor_party_name,
                    "[DEBTOR_STREET_NAME]": debtor_street_name if debtor_street_name != "" else "(No Address)",
                    "[DOB]": f", {dob}" if dob != "" else "",
                    "[ATTY_FIRST]": debtor_atty_first,
                    "[ATTY_LAST]": debtor_atty_last,
                    "[PARTY_ALT]": debtor_alt_names
                })
            all_debtors.append(case_debtors)
        else:
            debtor_party_name = get_debtor_values_from_debt(cases[i], 0, 'PARTY_NAME')
            debtor_street_name = get_debtor_values_from_debt(cases[i], 0, 'DEBTOR_STREET_NAME')
            dob = get_debtor_values_from_debt(cases[i], 0, 'DOB')
            debtor_atty_first = get_debtor_values_from_debt(cases[i], 0, 'ATTY_FIRST')
            debtor_atty_last = get_debtor_values_from_debt(cases[i], 0, 'ATTY_LAST')
            all_debtors.append([{
                "[PARTY_NAME]": debtor_party_name,
                "[DEBTOR_STREET_NAME]": debtor_street_name if debtor_street_name != "" else "",
                "[DOB]": f", {dob}" if dob != "" else "",
                "[ATTY_FIRST]": debtor_atty_first,
                "[ATTY_LAST]": debtor_atty_last,
                "[PARTY_ALT]": debtor_alt_names
            }])
            
                
    context = {
        "N_CASES": len(cases),
        "[case]": '',
        "[CLIENT_REF_NUM]": client_ref_num,
        "[JUDGMENT_NUMBER]": judgment_numbers,
        "[DATE_ENTERED]": dates_entered,
        "[DATE_SIGNED]": dates_signed, 
        "[CASE_NUM]": case_numbers,
        "[VENUE]": venues,
        "[ACTION_TYPE]": action_types,
        "[CASE_STATUS]": case_statuses,
        "[ORIG_DEBT_AMT]": orig_debt_amts,
        "[ALL_COMMENTS]": comments,
        "[DEBT_COMMENTS]": debt_comments,
        "CREDITORS": all_creditors,
        "DEBTORS": all_debtors,
        "[:forEach]": ''
    }
    
    document = Document(uri)    
    
    context["[ifDateSigned]Date Signed: [DATE_SIGNED][:if]"] = parse_date_signed(document, context)
    context["[ifComments]Comments: [COMMENTS][:if]"] = parse_state_comments(document, context)
    context["[ifDebtComment]Note: [DEBT_COMMENT][:if]"] = parse_state_debt_comments(document, context)
    
    
    
    newdoc = replace_docket_common(document, uri, 'CaseReport', context)
    newdoc.save(os.path.join(".",'jsnetwork_project','media',f'generated_CaseReport.docx'))
        

def get_element_value(element, tag_name):
    tag_elements = element.getElementsByTagName(tag_name)
    if tag_elements and tag_elements[0].firstChild:
        return tag_elements[0].firstChild.nodeValue
    return ''

def get_element_value_from_parties(element, tag_name):
    parties_elements = element.getElementsByTagName('PARTIES')
    if parties_elements and parties_elements[0].getElementsByTagName(tag_name):
        tag_elements = parties_elements[0].getElementsByTagName(tag_name)
        if tag_elements and tag_elements[0].firstChild.nodeValue:
            return tag_elements[0].firstChild.nodeValue
    return '#N/A' if tag_name in ['ATTORNEY','TRUSTEE'] else ''

def get_element_values_from_debt(element, tag_name):
    debt_element = element.getElementsByTagName('DEBT')
    if debt_element and debt_element[0].getElementsByTagName(tag_name):
        tag_elements = debt_element[0].getElementsByTagName(tag_name)
        if tag_elements and tag_elements[0].firstChild:
            return tag_elements[0].firstChild.nodeValue
    return ''

def get_debtor_values_from_debt(element, index, tag_name):
    debt_element = element.getElementsByTagName('DEBT')
    if debt_element and debt_element[0].getElementsByTagName('DEBTOR'):
        debtor_element = debt_element[0].getElementsByTagName('DEBTOR')[index]
        tag_elements = debtor_element.getElementsByTagName(tag_name)
        if tag_elements and tag_elements[0].firstChild:
            return tag_elements[0].firstChild.nodeValue
    return ''

def get_creditor_values_from_debt(element, index, tag_name):
    debt_element = element.getElementsByTagName('DEBT')
    if debt_element and debt_element[0].getElementsByTagName('CREDITOR'):
        creditor_element = debt_element[0].getElementsByTagName('CREDITOR')[index]
        tag_elements = creditor_element.getElementsByTagName(tag_name)
        if tag_elements and tag_elements[0].firstChild:
            return tag_elements[0].firstChild.nodeValue
    return ''
        

doc_uri = os.path.join(".", "jsnetwork_project", 'media', 'Case Report Template.docx')
xml = open(os.path.join(".", "jsnetwork_project", "media", "tom state.xml"), 'rb')
xml_data_str = xml.read().decode("utf-8")
xml_data = minidom.parseString(xml_data_str)
replace_state(doc_uri, xml_data)
# replace_usdc(doc_uri, xml_data)
# replace_docketreport(doc_uri, xml_data)
# replace_coverpage(doc_uri, xml_data)
# replace_patriot(doc_uri, xml_data)